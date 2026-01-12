import io
import re
import zipfile
from datetime import datetime
from typing import Dict, Any, List, Optional

import streamlit as st
from docx import Document
from openai import OpenAI
from pydantic import BaseModel, Field


# -----------------------------
# Helpers
# -----------------------------
def load_bytes(path: str) -> bytes:
    with open(path, "rb") as f:
        return f.read()


def load_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def extract_about_me_from_cv(base_cv: str) -> str:
    """
    Extract the 'about me' section from base_cv.txt
    Looks for common patterns like:
    - ABOUT ME:
    - About Me:
    - PROFESSIONAL SUMMARY:
    - Professional Summary:
    - PROFILE:
    Returns the content until the next section heading or empty if not found.
    """
    # Common section headers that might indicate the about me section
    about_me_patterns = [
        r"(?i)^ABOUT\s+ME[:\s]*$",
        r"(?i)^About\s+Me[:\s]*$",
        r"(?i)^PROFESSIONAL\s+SUMMARY[:\s]*$",
        r"(?i)^Professional\s+Summary[:\s]*$",
        r"(?i)^PROFILE[:\s]*$",
        r"(?i)^Profile[:\s]*$",
        r"(?i)^SUMMARY[:\s]*$",
        r"(?i)^Summary[:\s]*$",
    ]
    
    # Common next section headers to know where about me ends
    next_section_patterns = [
        r"(?i)^EXPERIENCE[:\s]*$",
        r"(?i)^WORK\s+EXPERIENCE[:\s]*$",
        r"(?i)^PROFESSIONAL\s+EXPERIENCE[:\s]*$",
        r"(?i)^EMPLOYMENT[:\s]*$",
        r"(?i)^EDUCATION[:\s]*$",
        r"(?i)^SKILLS[:\s]*$",
        r"(?i)^TECHNICAL\s+SKILLS[:\s]*$",
    ]
    
    lines = base_cv.split('\n')
    about_me_content = []
    found_start = False
    
    for i, line in enumerate(lines):
        # Check if this line is an "about me" header
        if not found_start:
            for pattern in about_me_patterns:
                if re.match(pattern, line.strip()):
                    found_start = True
                    break
            continue
        
        # If we found the start, collect content until next section
        stripped = line.strip()
        
        # Check if we hit the next section
        is_next_section = False
        for pattern in next_section_patterns:
            if re.match(pattern, stripped):
                is_next_section = True
                break
        
        if is_next_section:
            break
        
        # Add non-empty lines to about me content
        if stripped:
            about_me_content.append(stripped)
    
    result = ' '.join(about_me_content).strip()
    
    # If nothing found, return a helpful placeholder
    if not result:
        result = "Professional with extensive experience. [Note: Could not auto-extract 'about me' from base_cv.txt - please add a clear ABOUT ME: or PROFESSIONAL SUMMARY: section]"
    
    return result


def safe_filename(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "_", (s or "").strip())
    return s.strip("_") or "output"


def replace_tokens_in_doc(doc: Document, mapping: Dict[str, str]) -> None:
    """
    Replace {{TOKEN}} in paragraphs AND tables.
    Best practice: keep placeholders in their own paragraph or typed in one go.
    """
    def replace_in_paragraph(paragraph):
        if "{{" not in paragraph.text:
            return
        for run in paragraph.runs:
            for k, v in mapping.items():
                token = f"{{{{{k}}}}}"
                if token in run.text:
                    run.text = run.text.replace(token, v)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


def doc_contains_token(doc: Document, token: str) -> bool:
    """Check if a {{TOKEN}} exists anywhere in the document."""
    needle = f"{{{{{token}}}}}"

    for p in doc.paragraphs:
        if needle in p.text:
            return True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if needle in p.text:
                        return True

    return False


def make_zip(cv_bytes: bytes, cl_bytes: bytes, cv_name: str, cl_name: str) -> bytes:
    """Create an in-memory ZIP containing the CV and cover letter."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr(cv_name, cv_bytes)
        z.writestr(cl_name, cl_bytes)
    return buf.getvalue()


# -----------------------------
# Template maps (DE / NL × Corporate / Startup)
# -----------------------------
CV_TEMPLATES = {
    "Germany": {
        "Corporate": "templates/cv_template_DE_corporate.docx",
        "Startup": "templates/cv_template_DE_startup.docx",
    },
    "Netherlands": {
        "Corporate": "templates/cv_template_NL_corporate.docx",
        "Startup": "templates/cv_template_NL_startup.docx",
    },
}

CL_TEMPLATES = {
    "Germany": "templates/cover_letter_template_DE.docx",
    "Netherlands": "templates/cover_letter_template_NL.docx",
}


# -----------------------------
# Model output schema (Pydantic)
# 6 Versuni bullets = room for an extra bullet and better alignment.
# -----------------------------
class CVCoverOutput(BaseModel):
    company: str
    role_title: str
    about_me: str
    versuni_bullets: List[str] = Field(min_length=6, max_length=6)
    philips_bullets: List[str] = Field(min_length=2, max_length=2)
    cover_letter_body: str
    new_claims_not_in_base: List[str] = Field(default_factory=list)


# -----------------------------
# OpenAI call (Responses API + parse)
# -----------------------------
def generate_structured(
    job_description: str,
    prompt_rules: str,
    base_cv: str,
    base_cover_letter: str,
    base_about_me: str,
    model: str,
    country: str,
    cv_style: str,
    fix_mode: bool = False,
    issues_to_fix: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """
    fix_mode=False: normal generation
    fix_mode=True: second pass if the model flagged unsupported claims.
                   It must remove/adjust those claims AND still return 6 Versuni bullets.
    """
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

    issues_text = ""
    if fix_mode and issues_to_fix:
        issues_text = "\n".join([f"- {x}" for x in issues_to_fix])

    system_instructions = f"""
You are a CV + cover letter customization engine with a focus on hiring manager perspective.

TARGET CONTEXT
- Country: {country}
- CV positioning: {cv_style} role in {country}
- Make sure tone, examples and emphasis fit this market and style
  (e.g. for Germany vs Netherlands, and for corporate vs startup).

You are given:
1) The user's ORIGINAL CV content (base CV)
2) The user's ORIGINAL cover letter (base cover letter)
3) The user's ORIGINAL "about me" section (base about me)
4) A job description

OVERALL GOAL
- Tailor the CV and cover letter so that, when a recruiter scans them,
  the Versuni experience bullets and the cover letter bullet points
  clearly mirror the structure and priorities of the job description.

ABOUT ME SECTION - HIRING MANAGER REQUIREMENTS (CRITICAL):
As a hiring manager, the "about me" section is the FIRST thing I read. It must:

1. PRESERVE ALL FACTS AND NUMBERS:
   - Keep ALL years of experience, metrics, percentages, achievement numbers EXACTLY as stated in base_about_me
   - Do NOT change any numerical values (e.g., if it says "5 years", keep "5 years")
   - Do NOT modify any specific achievements or quantifiable results
   - Do NOT invent new numbers or metrics not present in the original

2. STRATEGIC ALIGNMENT WITH JOB DESCRIPTION:
   - Extract the top 3-5 key requirements from the job description
   - Reorder and emphasize content from base_about_me to prioritize what matches these requirements
   - Use similar terminology and keywords from the job description while keeping original facts
   - Lead with the most relevant experience for THIS specific role

3. STRUCTURE (3-5 sentences, ~80-120 words):
   - Opening sentence: Professional identity + years of experience (from base_about_me)
   - Middle 2-3 sentences: Key achievements with original numbers that align to job requirements
   - Closing sentence: Value proposition or what you're seeking (adapted to this role)

4. WRITING STYLE:
   - Use strong action verbs that appear in the job description
   - Mirror the tone (formal for corporate, dynamic for startup)
   - Keep the candidate's authentic voice from base_about_me
   - Make every word count - no filler phrases

5. WHAT TO AVOID:
   - Generic statements that could apply to anyone
   - Inventing experience or numbers not in base_about_me
   - Changing the order of facts unless it improves alignment with JD
   - Losing the candidate's personality or unique strengths

ALIGNMENT WITH JOB DESCRIPTION BULLET POINTS
- First, read the job description and mentally identify the main bullets
  under responsibilities / requirements / what you'll do / who you are.
- Keep their original order from top to bottom.
- For the Versuni section:
  - You MUST produce exactly 6 bullets.
  - Order these 6 bullets to follow the job description bullet order.
    Bullet 1 should align to the first key requirement, bullet 2 to the next, etc.
  - Each bullet should clearly show "I did this / I can do this" for that specific requirement,
    using matching concepts and key phrases where honest.
- For the cover letter body:
  - Include a clear section with bullet points (or short structured lines) that
    explicitly map your experience to the job requirements.
  - Order those cover letter bullets in the same sequence as the job description bullets,
    so a recruiter instantly sees the one-to-one match.

GROUNDING RULES (NO INVENTED CLAIMS)
- All content MUST be strictly grounded in the base CV + base cover letter + base about me.
- Do NOT invent employers, titles, dates, locations, tools, budgets, metrics, languages,
  or results that are not clearly supported by the base materials.
- You MAY rephrase and recombine what's already there to increase relevance and clarity.

WRITING STYLE
- Keep writing crisp, structured, and ATS-friendly.
- Use concrete outcomes and responsibilities where they genuinely exist in the base CV.
- Preserve the user's voice and tone from the base cover letter.

SCHEMA & SAFETY
- You must output exactly:
  - company
  - role_title
  - about_me (tailored from base_about_me with original facts preserved)
  - versuni_bullets (exactly 6)
  - philips_bullets (exactly 2)
  - cover_letter_body
  - new_claims_not_in_base (list of strings)
- IMPORTANT SAFETY CHECK:
  If you add any claim not clearly supported by the base CV, base cover letter, or base about me,
  list it in new_claims_not_in_base. Otherwise return [].

SECOND-PASS FIX MODE (only if requested):
- If fix_mode is enabled, you must eliminate or reword the unsupported claims listed below
  so that new_claims_not_in_base becomes [].
- You must still return 6 Versuni bullets and keep the JD-aligned ordering.

Unsupported claims to fix (if any):
{issues_text if issues_text else "(none)"}

BASE CV (source of truth):
{base_cv}

BASE COVER LETTER (source of truth):
{base_cover_letter}

BASE ABOUT ME (source of truth - preserve all facts and numbers):
{base_about_me}

USER PROMPT / RULES (additional preferences):
{prompt_rules}
""".strip()

    user_msg = f"JOB DESCRIPTION:\n{job_description}"
    if fix_mode:
        user_msg += (
            "\n\nPlease revise the draft to remove unsupported claims and keep "
            "new_claims_not_in_base empty, while maintaining strong impact, "
            "keeping 6 Versuni bullets aligned to the JD bullet order, "
            "preserving all original facts and numbers in about_me, and "
            "maintaining the cover letter structure."
        )

    resp = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": system_instructions},
            {"role": "user", "content": user_msg},
        ],
        text_format=CVCoverOutput,
    )

    return resp.output_parsed.model_dump()


# -----------------------------
# Build docs from templates
# -----------------------------
def build_docs(
    data: Dict[str, Any],
    cv_template_bytes: bytes,
    cl_template_bytes: bytes,
) -> tuple[bytes, bytes]:
    cv_doc = Document(io.BytesIO(cv_template_bytes))
    cl_doc = Document(io.BytesIO(cl_template_bytes))

    # Check whether the CV template has a dedicated slot for bullet 6.
    has_bullet_6 = doc_contains_token(cv_doc, "VERSUNI_BULLET_6")

    # If no placeholder for bullet 6, append it to bullet 5 so you still get the "extra bullet".
    bullet5 = data["versuni_bullets"][4]
    bullet6 = data["versuni_bullets"][5]
    if not has_bullet_6 and bullet6.strip():
        bullet5 = f"{bullet5}\n• {bullet6}"

    mapping = {
        "ROLE_TITLE": data["role_title"],
        "COMPANY": data["company"],
        "ABOUT_ME": data["about_me"],
        "VERSUNI_BULLET_1": data["versuni_bullets"][0],
        "VERSUNI_BULLET_2": data["versuni_bullets"][1],
        "VERSUNI_BULLET_3": data["versuni_bullets"][2],
        "VERSUNI_BULLET_4": data["versuni_bullets"][3],
        "VERSUNI_BULLET_5": bullet5,
        "VERSUNI_BULLET_6": bullet6 if has_bullet_6 else "",
        "PHILIPS_BULLET_1": data["philips_bullets"][0],
        "PHILIPS_BULLET_2": data["philips_bullets"][1],
        # You can adjust the city if you want to be country-specific.
        "DATE_TODAY": datetime.now().strftime("%d %B %Y, Berlin"),
        "COVER_LETTER_BODY": data["cover_letter_body"],
    }

    # Fill both documents
    replace_tokens_in_doc(cv_doc, mapping)
    replace_tokens_in_doc(cl_doc, mapping)

    # --- Remove duplicate "Cover Letter for ..." line between header and body ---
    # Keep the FIRST occurrence (title), blank out the SECOND occurrence if present.
    seen_title = False
    for p in cl_doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if "Cover Letter for" in text:
            if not seen_title:
                seen_title = True
            else:
                p.text = ""
                break

    cv_out = io.BytesIO()
    cl_out = io.BytesIO()
    cv_doc.save(cv_out)
    cl_doc.save(cl_out)
    return cv_out.getvalue(), cl_out.getvalue()


# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="CV + Cover Letter Agent", layout="wide")
st.title("CV + Cover Letter Agent")

model = st.secrets.get("MODEL", "gpt-5.2")

# -----------------------------
# Controls (country, style, debug, test JD)
# -----------------------------
colA, colB, colC = st.columns([1, 1, 1])
with colA:
    country = st.selectbox("Country", ["Germany", "Netherlands"], index=0)
with colB:
    cv_style = st.radio("CV style", ["Corporate", "Startup"], index=0, horizontal=True)
with colC:
    show_debug = st.checkbox("Show debug JSON", value=False)

use_test = st.checkbox("Use test job description (prompt/test_job_description.txt)", value=False)

# Resolve template paths from maps
cv_template_path = CV_TEMPLATES[country][cv_style]
cl_template_path = CL_TEMPLATES[country]

cv_template_bytes = load_bytes(cv_template_path)
cl_template_bytes = load_bytes(cl_template_path)

# Base materials
prompt_rules = load_text("prompt/instructions.txt")
base_cv = load_text("prompt/base_cv.txt")
base_cover_letter = load_text("prompt/base_cover_letter.txt")

# Auto-extract about me from base_cv
base_about_me = extract_about_me_from_cv(base_cv)

# Show extraction status
if "[Note: Could not auto-extract" in base_about_me:
    st.warning("⚠️ Could not automatically find 'About Me' section in base_cv.txt. Please ensure your CV has a clear section header like 'ABOUT ME:', 'PROFESSIONAL SUMMARY:', or 'PROFILE:'")
else:
    with st.expander("ℹ️ Auto-extracted 'About Me' section from base_cv.txt"):
        st.write(base_about_me)
        st.caption("This will be tailored for each job while preserving all facts and numbers.")

# Optional test JD
test_jd_path = "prompt/test_job_description.txt"
try:
    test_jd = load_text(test_jd_path)
except FileNotFoundError:
    test_jd = ""

st.caption("Choose country & style → paste a job description → generate → download ZIP with CV + cover letter (.docx).")
st.info("💡 **Enhanced:** The 'about me' section is now intelligently tailored to each job while preserving all your original facts and numbers!")

default_text = test_jd if (use_test and test_jd.strip()) else ""
job_description = st.text_area("Job description", value=default_text, height=320)

if st.button("Generate", type="primary"):
    if not job_description.strip():
        st.error("Please paste a job description.")
        st.stop()

    # Placeholder diagnostics
    cv_doc_check = Document(io.BytesIO(cv_template_bytes))
    cl_doc_check = Document(io.BytesIO(cl_template_bytes))

    required_cv_tokens = [
        "ABOUT_ME",
        "VERSUNI_BULLET_1", "VERSUNI_BULLET_2", "VERSUNI_BULLET_3", "VERSUNI_BULLET_4", "VERSUNI_BULLET_5",
        # Bullet 6 is optional in template.
        "PHILIPS_BULLET_1", "PHILIPS_BULLET_2",
    ]
    required_cl_tokens = ["COMPANY", "COVER_LETTER_BODY"]

    missing = []
    for t in required_cv_tokens:
        if not doc_contains_token(cv_doc_check, t):
            missing.append(f"CV template missing {{{{{t}}}}}")
    for t in required_cl_tokens:
        if not doc_contains_token(cl_doc_check, t):
            missing.append(f"Cover letter template missing {{{{{t}}}}}")

    if missing:
        st.error("Your templates are missing required placeholders:")
        for m in missing:
            st.write(f"- {m}")
        st.stop()

    # 1) First pass
    with st.spinner("Generating tailored CV with optimized 'about me' section..."):
        data = generate_structured(
            job_description=job_description,
            prompt_rules=prompt_rules,
            base_cv=base_cv,
            base_cover_letter=base_cover_letter,
            base_about_me=base_about_me,
            model=model,
            country=country,
            cv_style=cv_style,
            fix_mode=False,
        )

    # 2) If unsupported claims exist, auto-fix once
    if data.get("new_claims_not_in_base"):
        with st.spinner("Fixing unsupported claims and strengthening CV..."):
            data_fixed = generate_structured(
                job_description=job_description,
                prompt_rules=prompt_rules,
                base_cv=base_cv,
                base_cover_letter=base_cover_letter,
                base_about_me=base_about_me,
                model=model,
                country=country,
                cv_style=cv_style,
                fix_mode=True,
                issues_to_fix=data["new_claims_not_in_base"],
            )

        if data_fixed.get("new_claims_not_in_base"):
            st.error("I tried to fix unsupported claims, but some remain. Please review:")
            st.write(data_fixed["new_claims_not_in_base"])
            if show_debug:
                st.subheader("Debug: AI JSON output (after fix attempt)")
                st.json(data_fixed)
            st.stop()

        data = data_fixed

    # Build docs + ZIP
    cv_bytes, cl_bytes = build_docs(data, cv_template_bytes, cl_template_bytes)

    company = safe_filename(data["company"])
    role = safe_filename(data["role_title"])

    cv_name = f"Xiaoxuan_Li_CV_{company}_{role}.docx"
    cl_name = f"Xiaoxuan_Li_Cover_Letter_{company}_{role}.docx"
    zip_name = f"Xiaoxuan_Li_{company}_{role}.zip"

    zip_bytes = make_zip(cv_bytes, cl_bytes, cv_name, cl_name)

    st.success("Done! Your 'about me' section has been strategically aligned to this role.")
    
    # Preview the generated about me
    with st.expander("📝 Preview: Generated 'About Me' Section"):
        st.write(data["about_me"])
        st.caption("Compare this with your original to see how it's been tailored while preserving all facts.")
    
    st.download_button(
        "Download CV + Cover Letter (ZIP)",
        data=zip_bytes,
        file_name=zip_name,
        mime="application/zip",
    )

    if show_debug:
        st.subheader("Debug: AI JSON output")
        st.json(data)